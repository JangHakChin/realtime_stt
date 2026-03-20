
"""
AI Agent 백엔드 서버 (FastAPI)
- GPT-4o Function Calling으로 사용자 의도 파악
- 각 Tool을 자동 실행하고 결과 반환

필요 패키지 설치:
    pip install fastapi uvicorn openai python-docx sounddevice numpy scipy

실행 방법:
    python server.py
    → http://localhost:8000 에서 실행
"""

import os
import io
import json
import wave
import time
import queue
import threading
import tempfile
import datetime
import sys
from typing import Optional

import numpy as np
import sounddevice as sd
from fastapi import FastAPI
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from openai import OpenAI

# ─────────────────────────────────────────
# 설정
# ─────────────────────────────────────────
API_KEY  = os.getenv("OPENAI_API_KEY", "")
client   = OpenAI(api_key=API_KEY)
app      = FastAPI(title="AI Agent Server")

# CORS (브라우저에서 직접 호출 허용)
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# ─────────────────────────────────────────
# 요청/응답 모델
# ─────────────────────────────────────────
class ChatRequest(BaseModel):
    message: str
    context: Optional[str] = None   # 이전 대화 등 추가 컨텍스트

class ToolCall(BaseModel):
    name: str
    description: str
    status: str   # "running" | "done" | "error"

class ChatResponse(BaseModel):
    reply: str
    tool_calls: list[ToolCall] = []

class RecordingStatus(BaseModel):
    is_recording: bool
    duration: float        # 현재 녹음 경과 시간 (초)
    line_count: int        # 전사된 문장 수
    minutes_path: str      # 생성된 회의록 경로 (완료 시)


# ─────────────────────────────────────────
# 녹음 상태 관리 (서버 전역)
# ─────────────────────────────────────────
SAMPLE_RATE       = 16000
CHANNELS          = 1
CHUNK_DURATION    = 0.1
SILENCE_DURATION  = 1.2
SILENCE_THRESHOLD = 0.02
MIN_SPEECH_DUR    = 0.8
MAX_CHUNK_DUR     = 8.0
HALLUCINATION_PATTERNS = [
    "mbc","kbs","sbs","뉴스","이덕영","앵커",
    "시청해주셔서 감사합니다","구독과 좋아요","구독","좋아요","알림설정",
    "감사합니다","안녕하세요","bye","thank you for watching",
    "자막 제공","번역 제공","subtitles","copyright",
]

recording_state = {
    "is_recording":  False,
    "thread":        None,
    "stream":        None,
    "audio_queue":   None,
    "transcript":    [],        # 전사된 문장들
    "start_time":    None,
    "minutes_path":  "",
}


def _rms(chunk: np.ndarray) -> float:
    return float(np.sqrt(np.mean(chunk.astype(np.float32) ** 2))) / 32768

def _is_hallucination(text: str) -> bool:
    t = text.lower().strip()
    if len(t) <= 2: return True
    return any(p in t for p in HALLUCINATION_PATTERNS)

def _transcribe_chunk(audio_data: np.ndarray, language="ko") -> str:
    with io.BytesIO() as buf:
        with wave.open(buf, 'wb') as wf:
            wf.setnchannels(CHANNELS)
            wf.setsampwidth(2)
            wf.setframerate(SAMPLE_RATE)
            wf.writeframes(audio_data.tobytes())
        buf.seek(0)
        wav_bytes = buf.read()
    with tempfile.NamedTemporaryFile(suffix=".wav", delete=False) as tmp:
        tmp.write(wav_bytes); tmp_path = tmp.name
    try:
        with open(tmp_path, "rb") as f:
            result = client.audio.transcriptions.create(
                model="whisper-1", file=f, language=language)
        return result.text.strip()
    finally:
        os.unlink(tmp_path)

def _recording_loop(audio_queue: queue.Queue, language: str):
    """백그라운드 녹음 + 전사 루프"""
    speech_buffer, is_speaking = [], False
    silence_start, speech_start = None, None

    while recording_state["is_recording"]:
        try:
            chunk = audio_queue.get(timeout=0.5)
        except queue.Empty:
            continue

        volume    = _rms(chunk)
        is_speech = volume > SILENCE_THRESHOLD

        if is_speech:
            if not is_speaking:
                is_speaking  = True
                speech_start = time.time()
            silence_start = None
            speech_buffer.append(chunk)
        else:
            if is_speaking:
                speech_buffer.append(chunk)
                if silence_start is None:
                    silence_start = time.time()
                if time.time() - silence_start >= SILENCE_DURATION:
                    _flush(speech_buffer, language)
                    speech_buffer, is_speaking = [], False
                    silence_start, speech_start = None, None

        if is_speaking and speech_start and time.time() - speech_start >= MAX_CHUNK_DUR:
            _flush(speech_buffer, language)
            speech_buffer = []
            speech_start  = time.time()
            silence_start = None

    # 루프 종료 → 남은 버퍼 처리
    if speech_buffer:
        _flush(speech_buffer, language)

def _flush(speech_buffer: list, language: str):
    audio = np.concatenate(speech_buffer, axis=0).flatten()
    if len(audio) / SAMPLE_RATE < MIN_SPEECH_DUR:
        return
    try:
        text = _transcribe_chunk(audio, language)
        if text and not _is_hallucination(text):
            recording_state["transcript"].append(text)
    except Exception:
        pass

def _generate_minutes() -> str:
    """GPT-4o로 회의록 생성 후 docx 저장, 경로 반환"""
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    transcript = recording_state["transcript"]
    if not transcript:
        return ""

    start_time = recording_state["start_time"]
    transcript_text = "\n".join(transcript)

    # GPT-4o로 구조화
    prompt = f"""다음은 회의 전사 내용이야.
JSON 형식으로 회의록을 작성해줘:
{{
  "title": "회의 제목",
  "agenda": "안건 한 줄 요약",
  "discussions": [{{"topic": "주제", "content": "내용"}}],
  "decisions": ["결정 사항"],
  "action_items": [{{"task": "할 일", "owner": "담당자"}}]
}}

전사 내용:
{transcript_text}"""

    res = client.chat.completions.create(
        model="gpt-4o",
        messages=[{"role": "user", "content": prompt}],
        response_format={"type": "json_object"},
    )
    data = json.loads(res.choices[0].message.content)

    # docx 생성
    doc = Document()
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Cm(2.5)
        section.left_margin = section.right_margin = Cm(3.0)

    def sf(run, size=11, bold=False, color=None):
        run.font.name = "맑은 고딕"; run.font.size = Pt(size); run.font.bold = bold
        if color: run.font.color.rgb = RGBColor(*color)
        rPr = run._r.get_or_add_rPr()
        rF = OxmlElement('w:rFonts'); rF.set(qn('w:eastAsia'), "맑은 고딕"); rPr.insert(0, rF)

    def heading(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(4)
        pPr = p._p.get_or_add_pPr(); pBdr = OxmlElement('w:pBdr')
        bot = OxmlElement('w:bottom'); bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),'6')
        bot.set(qn('w:space'),'1'); bot.set(qn('w:color'),'1F497D'); pBdr.append(bot); pPr.append(pBdr)
        r = p.add_run(text); sf(r, size=13, bold=True, color=(0x1F,0x49,0x7D))

    def bullet(text):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(2)
        r = p.add_run(text); sf(r, size=10.5)

    # 제목
    tp = doc.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = tp.add_run("회  의  록"); sf(r, size=22, bold=True, color=(0x1F,0x49,0x7D))
    sp = doc.add_paragraph(); sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sp.paragraph_format.space_after = Pt(16)
    r2 = sp.add_run(data.get("title","회의")); sf(r2, size=12, color=(0x40,0x40,0x40))

    heading("1. 주요 논의 사항")
    for item in data.get("discussions", []):
        p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(4)
        r = p.add_run(f"▶ {item['topic']}"); sf(r, size=10.5, bold=True, color=(0x1F,0x49,0x7D))
        bullet(item["content"])

    heading("2. 결정 사항")
    for d in (data.get("decisions") or ["특별한 결정 사항 없음"]):
        bullet(d)

    heading("3. 후속 조치")
    for ai in (data.get("action_items") or [{"task":"없음","owner":""}]):
        bullet(f"{ai['task']}  {'— ' + ai['owner'] if ai.get('owner') else ''}")

    # 푸터
    doc.add_paragraph().paragraph_format.space_after = Pt(20)
    fp = doc.add_paragraph(); fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = fp.add_run("※ 본 회의록은 STT + GPT-4o Agent가 자동 생성하였습니다.")
    sf(r, size=9, color=(0x80,0x80,0x80))

    fname = f"회의록_{start_time.strftime('%Y%m%d_%H%M')}.docx"
    out   = os.path.join(os.path.dirname(__file__), fname)
    doc.save(out)
    return os.path.abspath(out)


# ─────────────────────────────────────────
# Tool 정의 (GPT-4o에게 알려줄 도구 목록)
# ─────────────────────────────────────────
TOOLS = [
    {
        "type": "function",
        "function": {
            "name": "start_meeting_recorder",
            "description": "실시간으로 회의를 녹음하고 STT로 전사한 뒤, 종료 시 자동으로 회의록 docx 파일을 생성한다. '회의 녹음', '회의록 만들어줘', 'STT' 등의 요청에 사용한다.",
            "parameters": {
                "type": "object",
                "properties": {
                    "language": {
                        "type": "string",
                        "description": "녹음할 언어 코드 (기본값: ko)",
                        "enum": ["ko", "en", "ja", "zh"]
                    }
                },
                "required": []
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "translate_text",
            "description": "주어진 텍스트를 지정한 언어로 번역한다.",
            "parameters": {
                "type": "object",
                "properties": {
                    "text":        {"type": "string", "description": "번역할 텍스트"},
                    "target_lang": {"type": "string", "description": "번역 대상 언어 (예: 영어, 일본어, 중국어)"}
                },
                "required": ["text", "target_lang"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "summarize_conversation",
            "description": "지금까지 나눈 대화 내용을 요약한다.",
            "parameters": {
                "type": "object",
                "properties": {
                    "conversation": {"type": "string", "description": "요약할 대화 내용"}
                },
                "required": ["conversation"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "create_document",
            "description": "주어진 내용을 바탕으로 Word 문서(docx) 또는 PPT를 생성한다.",
            "parameters": {
                "type": "object",
                "properties": {
                    "content":  {"type": "string", "description": "문서에 담을 내용"},
                    "doc_type": {
                        "type": "string",
                        "description": "문서 종류",
                        "enum": ["docx", "pptx"]
                    }
                },
                "required": ["content", "doc_type"]
            }
        }
    }
]


# ─────────────────────────────────────────
# Tool 실행 함수들
# ─────────────────────────────────────────
def run_start_meeting_recorder(language: str = "ko") -> str:
    """서버 내부에서 녹음 시작 (브라우저 버튼으로 제어 가능)"""
    if recording_state["is_recording"]:
        return "이미 녹음이 진행 중입니다."
    recording_state["is_recording"] = True
    recording_state["transcript"]   = []
    recording_state["start_time"]   = datetime.datetime.now()
    recording_state["minutes_path"] = ""

    aq = queue.Queue()
    recording_state["audio_queue"] = aq

    def callback(indata, frames, t, status):
        aq.put(indata.copy())

    stream = sd.InputStream(
        samplerate=SAMPLE_RATE, channels=CHANNELS,
        dtype='int16', blocksize=int(SAMPLE_RATE * CHUNK_DURATION),
        callback=callback,
    )
    recording_state["stream"] = stream
    stream.start()

    t = threading.Thread(
        target=_recording_loop, args=(aq, language), daemon=True)
    recording_state["thread"] = t
    t.start()

    return "RECORDING_STARTED"


def run_translate_text(text: str, target_lang: str) -> str:
    """GPT-4o로 텍스트 번역"""
    res = client.chat.completions.create(
        model="gpt-4o",
        messages=[{
            "role": "user",
            "content": f"다음 텍스트를 {target_lang}로 번역해줘. 번역 결과만 출력해.\n\n{text}"
        }]
    )
    return res.choices[0].message.content.strip()


def run_summarize_conversation(conversation: str) -> str:
    """GPT-4o로 대화 요약"""
    res = client.chat.completions.create(
        model="gpt-4o",
        messages=[{
            "role": "user",
            "content": f"다음 대화를 핵심만 3~5줄로 한국어로 요약해줘.\n\n{conversation}"
        }]
    )
    return res.choices[0].message.content.strip()


def run_create_document(content: str, doc_type: str) -> str:
    """문서 생성 (docx / pptx)"""
    if doc_type == "docx":
        return f"📄 Word 문서 생성 기능은 준비 중입니다. (내용: {content[:50]}...)"
    elif doc_type == "pptx":
        return f"📊 PPT 생성 기능은 준비 중입니다. (내용: {content[:50]}...)"
    return "지원하지 않는 문서 형식입니다."


# Tool 이름 → 실행 함수 매핑
TOOL_EXECUTORS = {
    "start_meeting_recorder": lambda args: run_start_meeting_recorder(**args),
    "translate_text":         lambda args: run_translate_text(**args),
    "summarize_conversation": lambda args: run_summarize_conversation(**args),
    "create_document":        lambda args: run_create_document(**args),
}

TOOL_DISPLAY_NAMES = {
    "start_meeting_recorder": "회의 녹음 & 회의록 생성",
    "translate_text":         "텍스트 번역",
    "summarize_conversation": "대화 요약",
    "create_document":        "문서 생성",
}


# ─────────────────────────────────────────
# API 엔드포인트
# ─────────────────────────────────────────
@app.get("/health")
async def health():
    return {"status": "ok"}


@app.post("/recording/start")
async def recording_start(language: str = "ko"):
    msg = run_start_meeting_recorder(language)
    return {"status": "started" if msg == "RECORDING_STARTED" else "error", "message": msg}


@app.post("/recording/stop")
async def recording_stop():
    if not recording_state["is_recording"]:
        return {"status": "error", "message": "녹음 중이 아닙니다."}

    recording_state["is_recording"] = False
    if recording_state["stream"]:
        recording_state["stream"].stop()
        recording_state["stream"].close()

    # 회의록 생성 (동기 처리)
    path = _generate_minutes()
    recording_state["minutes_path"] = path

    line_count = len(recording_state["transcript"])
    return {
        "status": "stopped",
        "line_count":   line_count,
        "minutes_path": path,
        "transcript":   recording_state["transcript"],
    }


@app.get("/recording/status", response_model=RecordingStatus)
async def recording_status():
    elapsed = 0.0
    if recording_state["is_recording"] and recording_state["start_time"]:
        elapsed = (datetime.datetime.now() - recording_state["start_time"]).total_seconds()
    return RecordingStatus(
        is_recording = recording_state["is_recording"],
        duration     = round(elapsed, 1),
        line_count   = len(recording_state["transcript"]),
        minutes_path = recording_state["minutes_path"],
    )


@app.post("/chat", response_model=ChatResponse)
async def chat(req: ChatRequest):
    messages = [
        {
            "role": "system",
            "content": (
                "너는 한국어를 사용하는 AI Assistant야. "
                "사용자의 요청을 파악해서 적절한 도구(tool)를 사용해. "
                "도구 실행 후 결과를 자연스럽게 한국어로 설명해줘. "
                "도구가 필요 없는 질문은 그냥 친절하게 답변해."
            )
        },
        {"role": "user", "content": req.message}
    ]

    tool_calls_result: list[ToolCall] = []

    # 1단계: GPT-4o에게 Tool 목록과 함께 요청
    response = client.chat.completions.create(
        model="gpt-4o",
        messages=messages,
        tools=TOOLS,
        tool_choice="auto",
    )

    msg = response.choices[0].message

    # 2단계: Tool 호출이 있으면 실행
    if msg.tool_calls:
        messages.append(msg)  # assistant 메시지 추가

        for tc in msg.tool_calls:
            fn_name = tc.function.name
            fn_args = json.loads(tc.function.arguments)
            display_name = TOOL_DISPLAY_NAMES.get(fn_name, fn_name)

            # Tool 실행
            try:
                executor = TOOL_EXECUTORS.get(fn_name)
                if executor:
                    result = executor(fn_args)
                    status = "done"
                else:
                    result = f"'{fn_name}' 도구를 찾을 수 없습니다."
                    status = "error"
            except Exception as e:
                result = f"오류: {e}"
                status = "error"

            tool_calls_result.append(ToolCall(
                name=display_name,
                description=result[:120],
                status=status,
            ))

            # Tool 결과를 messages에 추가
            messages.append({
                "role": "tool",
                "tool_call_id": tc.id,
                "content": result,
            })

        # 3단계: Tool 결과를 포함해 최종 응답 생성
        final = client.chat.completions.create(
            model="gpt-4o",
            messages=messages,
        )
        reply = final.choices[0].message.content.strip()

    else:
        # Tool 없이 바로 응답
        reply = msg.content.strip()

    return ChatResponse(reply=reply, tool_calls=tool_calls_result)


# ─────────────────────────────────────────
# 실행
# ─────────────────────────────────────────
if __name__ == "__main__":
    import uvicorn
    if API_KEY == "여기에_API_KEY_입력":
        print("⚠️  API 키를 설정해주세요!")
        exit(1)
    print("🚀 AI Agent 서버 시작: http://localhost:8000")
    print("   브라우저에서 index.html 을 열어주세요.")
    uvicorn.run(app, host="0.0.0.0", port=8000)