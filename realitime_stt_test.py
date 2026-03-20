"""
실시간 STT (Speech-to-Text) + 자동 회의록 생성
- 마이크 입력을 실시간으로 감지
- 말이 끝난 구간을 자동 감지 (VAD)
- OpenAI Whisper API로 텍스트 변환
- Ctrl+C 종료 시 GPT-4o가 자동으로 회의록 .docx 생성
 
필요 패키지 설치:
    pip install openai sounddevice numpy scipy python-docx
 
실행 방법:
    OPENAI_API_KEY=your_key_here python realtime_stt.py
    또는 스크립트 내 API_KEY 변수에 직접 입력
"""
 
import os
import io
import time
import datetime
import threading
import queue
import tempfile
import wave
 
import numpy as np
import sounddevice as sd
from openai import OpenAI
 
# ─────────────────────────────────────────
# 설정
# ─────────────────────────────────────────
API_KEY          = os.getenv("OPENAI_API_KEY", "")
LANGUAGE         = "ko"          # 언어 코드 (ko=한국어, en=영어, ja=일본어 등)
SAMPLE_RATE      = 16000         # 샘플레이트 (Whisper 권장: 16000)
CHANNELS         = 1             # 모노
CHUNK_DURATION   = 0.1           # 오디오 버퍼 단위 (초)
SILENCE_DURATION = 1.2           # 이 시간(초) 이상 조용하면 → 문장 완성으로 판단
SILENCE_THRESHOLD = 0.02         # 볼륨 임계값 (높일수록 민감도 낮아짐 / 환각 줄어듦)
MIN_SPEECH_DURATION = 0.8        # 최소 발화 길이 (초) - 너무 짧은 노이즈 제거
MAX_CHUNK_DURATION  = 8.0        # 최대 버퍼 길이 (초) - 말이 계속 이어져도 강제 전송
OUTPUT_DIR       = "."           # 회의록 저장 폴더 (기본: 현재 폴더)
 
# Whisper 환각 필터 - 자주 나오는 가짜 문구 목록
HALLUCINATION_PATTERNS = [
    "mbc", "kbs", "sbs", "뉴스", "이덕영", "앵커",
    "시청해주셔서 감사합니다", "구독과 좋아요", "구독", "좋아요", "알림설정",
    "감사합니다", "안녕하세요", "bye", "thank you for watching",
    "자막 제공", "번역 제공", "subtitles", "copyright",
]
 
 
# ─────────────────────────────────────────
# 회의록 Agent
# ─────────────────────────────────────────
class MinutesAgent:
    """전사된 텍스트를 받아 GPT-4o로 회의록을 생성하고 docx로 저장"""
 
    def __init__(self, client: OpenAI):
        self.client = client
 
    def generate(self, transcript_lines: list[str], start_time: datetime.datetime) -> str:
        """GPT-4o로 회의록 구조화"""
        transcript = "\n".join(transcript_lines)
        today = start_time.strftime("%Y년 %m월 %d일")
        start_str = start_time.strftime("%H:%M")
 
        prompt = f"""다음은 오늘({today}) 회의에서 실시간 STT로 전사된 내용이야.
이 내용을 바탕으로 아래 형식에 맞춰 한국어로 회의록을 작성해줘.
내용이 불분명하거나 잡음인 부분은 자연스럽게 정리하고, 중요한 논의만 포함해.
 
출력 형식 (JSON):
{{
  "title": "회의 제목 (내용 기반으로 추론)",
  "agenda": "회의 목적/안건 한 줄 요약",
  "discussions": [
    {{"topic": "주제명", "content": "상세 내용 1~3줄"}}
  ],
  "decisions": ["결정 사항 1", "결정 사항 2"],
  "action_items": [
    {{"task": "할 일", "owner": "담당자 (불명확하면 '미정')"}}
  ]
}}
 
전사 내용:
{transcript}
"""
        response = self.client.chat.completions.create(
            model="gpt-4o",
            messages=[{"role": "user", "content": prompt}],
            response_format={"type": "json_object"},
        )
        return response.choices[0].message.content
 
    def save_docx(self, json_str: str, start_time: datetime.datetime) -> str:
        """회의록 JSON → docx 파일로 저장"""
        import json
        from docx import Document
        from docx.shared import Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
 
        data = json.loads(json_str)
 
        doc = Document()
        for section in doc.sections:
            section.top_margin    = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin   = Cm(3.0)
            section.right_margin  = Cm(3.0)
 
        # ── 헬퍼 ──
        def set_font(run, size=11, bold=False, color=None):
            run.font.name = "맑은 고딕"
            run.font.size = Pt(size)
            run.font.bold = bold
            if color:
                run.font.color.rgb = RGBColor(*color)
            rPr = run._r.get_or_add_rPr()
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:eastAsia'), "맑은 고딕")
            rPr.insert(0, rFonts)
 
        def add_para(text="", bold=False, size=11, color=None, align=None, sb=0, sa=6):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(sb)
            p.paragraph_format.space_after  = Pt(sa)
            if align: p.alignment = align
            if text:
                r = p.add_run(text)
                set_font(r, size=size, bold=bold, color=color)
            return p
 
        def add_section(title):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after  = Pt(4)
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bot = OxmlElement('w:bottom')
            bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '6')
            bot.set(qn('w:space'), '1');    bot.set(qn('w:color'), '1F497D')
            pBdr.append(bot); pPr.append(pBdr)
            r = p.add_run(title)
            set_font(r, size=13, bold=True, color=(0x1F, 0x49, 0x7D))
 
        def add_bullet(text):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(2)
            r = p.add_run(text)
            set_font(r, size=10.5)
 
        def shade_cell(cell, hex_color):
            tcPr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), hex_color)
            tcPr.append(shd)
 
        def cell_border(cell):
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for side in ['top','left','bottom','right']:
                el = OxmlElement(f'w:{side}')
                el.set(qn('w:val'), 'single'); el.set(qn('w:sz'), '4')
                el.set(qn('w:color'), 'CCCCCC')
                tcBorders.append(el)
            tcPr.append(tcBorders)
 
        # ── 제목 ──
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run("회  의  록")
        set_font(r, size=22, bold=True, color=(0x1F, 0x49, 0x7D))
 
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(16)
        r2 = p2.add_run(data.get("title", "회의"))
        set_font(r2, size=12, color=(0x40, 0x40, 0x40))
 
        # ── 구분선 ──
        hr = doc.add_paragraph()
        hr.paragraph_format.space_after = Pt(12)
        pPr = hr._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bot = OxmlElement('w:bottom')
        bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '12')
        bot.set(qn('w:space'), '1');    bot.set(qn('w:color'), '1F497D')
        pBdr.append(bot); pPr.append(pBdr)
 
        # ── 기본 정보 테이블 ──
        tbl = doc.add_table(rows=2, cols=4)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        widths = [Cm(2.8), Cm(5.5), Cm(2.8), Cm(5.5)]
 
        def fill(row, col, text, label=False):
            cell = tbl.cell(row, col)
            cell.width = widths[col]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell_border(cell)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)
            r = p.add_run(f"  {text}")
            if label:
                set_font(r, size=10, bold=True, color=(255,255,255))
                shade_cell(cell, '1F497D')
            else:
                set_font(r, size=10)
                shade_cell(cell, 'F5F8FC')
 
        fill(0,0,"회의 일시", True); fill(0,1, start_time.strftime("%Y년 %m월 %d일  %H:%M"))
        fill(0,2,"작성 방법", True); fill(0,3,"STT 자동 전사 + GPT-4o 정리")
        fill(1,0,"안    건", True)
        c = tbl.cell(1,1); tbl.cell(1,1).merge(tbl.cell(1,3))
        c.paragraph_format if False else None
        cell_border(c); shade_cell(c, 'F5F8FC')
        p = c.paragraphs[0]
        p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(3)
        r = p.add_run(f"  {data.get('agenda','')}")
        set_font(r, size=10)
 
        doc.add_paragraph().paragraph_format.space_after = Pt(4)
 
        # ── 주요 논의 사항 ──
        add_section("1. 주요 논의 사항")
        for item in data.get("discussions", []):
            add_para(f"▶ {item['topic']}", bold=True, size=10.5,
                     color=(0x1F,0x49,0x7D), sb=6, sa=2)
            add_bullet(item["content"])
 
        # ── 결정 사항 ──
        add_section("2. 결정 사항")
        decisions = data.get("decisions", [])
        if decisions:
            for d in decisions:
                add_bullet(d)
        else:
            add_bullet("특별한 결정 사항 없음")
 
        # ── 후속 조치 ──
        add_section("3. 후속 조치 (Action Items)")
        action_items = data.get("action_items", [])
        if action_items:
            at = doc.add_table(rows=len(action_items)+1, cols=2)
            at.alignment = WD_TABLE_ALIGNMENT.CENTER
            for j, (txt, w) in enumerate(zip(["후속 조치 내용", "담당"],
                                              [Cm(11.5), Cm(3.1)])):
                cell = at.cell(0, j)
                cell.width = w; cell_border(cell); shade_cell(cell, '1F497D')
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(3)
                p.paragraph_format.space_after  = Pt(3)
                r = p.add_run(txt)
                set_font(r, size=10, bold=True, color=(255,255,255))
 
            for i, ai in enumerate(action_items):
                bg = 'F5F8FC' if i % 2 == 0 else 'FFFFFF'
                c0, c1 = at.cell(i+1, 0), at.cell(i+1, 1)
                for c in [c0, c1]:
                    cell_border(c); shade_cell(c, bg)
                    c.paragraphs[0].paragraph_format.space_before = Pt(3)
                    c.paragraphs[0].paragraph_format.space_after  = Pt(3)
                r0 = c0.paragraphs[0].add_run(ai["task"])
                set_font(r0, size=10)
                c1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                r1 = c1.paragraphs[0].add_run(ai["owner"])
                set_font(r1, size=10)
        else:
            add_bullet("후속 조치 없음")
 
        # ── 하단 ──
        doc.add_paragraph().paragraph_format.space_after = Pt(20)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p.add_run("※ 본 회의록은 실시간 STT + GPT-4o Agent가 자동 생성하였습니다.")
        set_font(r, size=9, color=(0x80, 0x80, 0x80))
 
        # ── 저장 ──
        filename = f"회의록_{start_time.strftime('%Y%m%d_%H%M')}.docx"
        out_path = os.path.join(OUTPUT_DIR, filename)
        doc.save(out_path)
        return out_path
 
 
# ─────────────────────────────────────────
# STT 클래스
# ─────────────────────────────────────────
class RealtimeSTT:
    def __init__(self):
        self.client = OpenAI(api_key=API_KEY)
        self.audio_queue = queue.Queue()
        self.is_running = False
 
        # 상태 추적
        self.speech_buffer = []
        self.is_speaking = False
        self.silence_start = None
        self.speech_start = None
 
        # 전사 누적 (회의록 용)
        self.transcript_lines: list[str] = []
        self.session_start = datetime.datetime.now()
 
        # Agent
        self.agent = MinutesAgent(self.client)
 
    def _rms(self, audio_chunk: np.ndarray) -> float:
        return float(np.sqrt(np.mean(audio_chunk.astype(np.float32) ** 2))) / 32768
 
    def _transcribe(self, audio_data: np.ndarray) -> str:
        with io.BytesIO() as wav_buffer:
            with wave.open(wav_buffer, 'wb') as wf:
                wf.setnchannels(CHANNELS)
                wf.setsampwidth(2)
                wf.setframerate(SAMPLE_RATE)
                wf.writeframes(audio_data.tobytes())
            wav_buffer.seek(0)
            wav_bytes = wav_buffer.read()
 
        with tempfile.NamedTemporaryFile(suffix=".wav", delete=False) as tmp:
            tmp.write(wav_bytes)
            tmp_path = tmp.name
 
        try:
            with open(tmp_path, "rb") as audio_file:
                result = self.client.audio.transcriptions.create(
                    model="whisper-1",
                    file=audio_file,
                    language=LANGUAGE,
                )
            return result.text.strip()
        finally:
            os.unlink(tmp_path)
 
    def _is_hallucination(self, text: str) -> bool:
        text_lower = text.lower().strip()
        if len(text_lower) <= 2:
            return True
        for pattern in HALLUCINATION_PATTERNS:
            if pattern in text_lower:
                return True
        return False
 
    def _flush_buffer(self, force=False):
        full_audio = np.concatenate(self.speech_buffer, axis=0).flatten()
        duration = len(full_audio) / SAMPLE_RATE
 
        if duration >= MIN_SPEECH_DURATION:
            print("(인식 중...)", end="\r", flush=True)
            try:
                text = self._transcribe(full_audio)
                if text and not self._is_hallucination(text):
                    print(f"\r💬 {text}          ")
                    self.transcript_lines.append(text)   # ← 누적
                else:
                    print("\r", end="", flush=True)
            except Exception as e:
                print(f"\r❌ 오류: {e}")
        else:
            print("\r", end="", flush=True)
 
        self.speech_buffer = []
        self.silence_start = None
        self.speech_start = time.time() if force else None
        if not force:
            self.is_speaking = False
 
    def _generate_minutes(self):
        """회의록 자동 생성 (종료 시 호출)"""
        if not self.transcript_lines:
            print("  전사된 내용이 없어 회의록을 생성하지 않습니다.")
            return
 
        print("\n📝 회의록 생성 중... (GPT-4o 분석 중)")
        try:
            json_str  = self.agent.generate(self.transcript_lines, self.session_start)
            out_path  = self.agent.save_docx(json_str, self.session_start)
            print(f"✅ 회의록 저장 완료: {os.path.abspath(out_path)}")
        except Exception as e:
            print(f"❌ 회의록 생성 실패: {e}")
 
    def _audio_callback(self, indata, frames, time_info, status):
        if status:
            pass
        self.audio_queue.put(indata.copy())
 
    def _process_loop(self):
        while self.is_running:
            try:
                chunk = self.audio_queue.get(timeout=0.5)
            except queue.Empty:
                continue
 
            volume = self._rms(chunk)
            is_speech = volume > SILENCE_THRESHOLD
 
            if is_speech:
                if not self.is_speaking:
                    self.is_speaking = True
                    self.speech_start = time.time()
                    print("\n🎤 ", end="", flush=True)
                self.silence_start = None
                self.speech_buffer.append(chunk)
            else:
                if self.is_speaking:
                    self.speech_buffer.append(chunk)
                    if self.silence_start is None:
                        self.silence_start = time.time()
                    if time.time() - self.silence_start >= SILENCE_DURATION:
                        self._flush_buffer()
 
            if self.is_speaking and self.speech_start is not None:
                if time.time() - self.speech_start >= MAX_CHUNK_DURATION:
                    print(" ✂️ (분할 전송)", end="", flush=True)
                    self._flush_buffer(force=True)
 
    def start(self):
        print("=" * 52)
        print("  실시간 STT + 자동 회의록 생성")
        print(f"  언어: {LANGUAGE} | 무음 감지: {SILENCE_DURATION}초 | 최대 청크: {MAX_CHUNK_DURATION}초")
        print("  말하면 자동으로 텍스트로 변환됩니다.")
        print("  종료(Ctrl+C) 시 회의록이 자동 생성됩니다.")
        print("=" * 52)
 
        self.is_running = True
        process_thread = threading.Thread(target=self._process_loop, daemon=True)
        process_thread.start()
 
        with sd.InputStream(
            samplerate=SAMPLE_RATE,
            channels=CHANNELS,
            dtype='int16',
            blocksize=int(SAMPLE_RATE * CHUNK_DURATION),
            callback=self._audio_callback,
        ):
            try:
                while True:
                    time.sleep(0.1)
            except KeyboardInterrupt:
                print("\n\n⏹ 녹음 종료.")
                self.is_running = False
                self._generate_minutes()
 
 
# ─────────────────────────────────────────
# 실행
# ─────────────────────────────────────────
if __name__ == "__main__":
    if API_KEY == "여기에_API_KEY_입력":
        print("⚠️  API 키를 설정해주세요!")
        print("   방법 1: 환경변수  →  export OPENAI_API_KEY=sk-...")
        print("   방법 2: 스크립트 상단 API_KEY 변수에 직접 입력")
        exit(1)
 
    stt = RealtimeSTT()
    stt.start()